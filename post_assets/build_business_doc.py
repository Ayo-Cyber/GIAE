"""Build GIAE_Business_Case.docx — a mentor-ready business memo with figures.

Mirrors BUSINESS.md, formatted as an editable Word document with the two
figures that carry the argument: gene-finding parity (commodity) and the
calibration before/after (the trust differentiator).

  PYTHONPATH=src .venv/bin/python post_assets/build_business_doc.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "GIAE_Business_Case.docx"

INDIGO = RGBColor(0x51, 0x57, 0xDD)
INK = RGBColor(0x1A, 0x20, 0x30)
MUTED = RGBColor(0x5B, 0x64, 0x80)
AMBER = RGBColor(0xB7, 0x79, 0x1F)
GREEN = RGBColor(0x0E, 0xA6, 0x72)

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)


def heading(text, color=INDIGO, size=15, space_before=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def kicker(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    return p


def body(runs):
    """runs: list of (text, {bold, italic, color})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for text, fmt in runs:
        r = p.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if "color" in fmt:
            r.font.color.rgb = fmt["color"]
    return p


def bullet(runs):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for text, fmt in runs:
        r = p.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if "color" in fmt:
            r.font.color.rgb = fmt["color"]


def figure(path: Path, caption: str, width=6.3):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


# ── Title ────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
tr = t.add_run("GIAE — Business Case")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = INK
sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(6)
sr = sub.add_run("Auditable, calibrated genome interpretation for high-stakes "
                 "decisions — starting with phage therapy.")
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = INDIGO
body([("We don't sell the algorithm. We sell the confidence to act on the answer.", {})])

# ── 1. Starting point ────────────────────────────────────────────────────────
kicker("The honest starting point")
heading("Gene annotation is free. Trust isn't.")
body([("The tension — the core is a commodity. ", {"bold": True, "color": AMBER}),
      ("Bakta, Prokka, PGAP and DFAST are free, open-source and excellent. Our "
       "own 35-genome benchmark showed GIAE's ", {}),
      ("gene finding is Prodigal", {"italic": True}),
      (" — the same engine everyone runs. You cannot sell “annotation”; "
       "the market price of the algorithm is $0.", {})])
body([("The answer — sell what free tools structurally can't. ", {"bold": True, "color": GREEN}),
      ("A free command-line tool hands you a label with no confidence and no "
       "reasoning. GIAE adds calibrated confidence (out-of-sample ECE 0.004), "
       "auditable provenance, and honest abstention. In high-stakes work, a "
       "confident wrong answer costs more than “I don't know.”", {})])

figure(ROOT / "post_assets" / "pitch_scatter_f1.png",
       "Figure 1 — GIAE vs Bakta gene-finding F1 across 35 genomes. Points near "
       "the diagonal mean parity: both tools call genes with Prodigal, so the "
       "algorithm is a commodity. This is exactly why we don't compete on "
       "annotation — and the setup for why trust is the product.")

# ── 2. Beachhead ─────────────────────────────────────────────────────────────
kicker("Beachhead")
heading("Win one vertical where trust is worth money")
body([("Academics won't pay for annotation they get free on their cluster — they "
       "are the credibility channel, not the customer. The wedge is a market "
       "that is funded, underserved by tooling, and where a wrong call has "
       "consequences. GIAE is already phage-aware, and that market exists today.", {})])
bullet([("Phage therapy (primary). ", {"bold": True}),
        ("Companies putting phages into patients need annotation plus safety "
         "screening — lysogeny, AMR, toxin and virulence genes — with an "
         "auditable, reproducible, regulatory-defensible report. Real pain, real "
         "budget, almost no good tooling.", {})])
bullet([("Biosecurity & biosurveillance (adjacent). ", {"bold": True}),
        ("Government and public-health screening for hazards and novel organisms. "
         "Calibrated, auditable calls are exactly what agencies require. Deeper "
         "pockets, longer procurement.", {})])

# ── 3. Who pays ──────────────────────────────────────────────────────────────
kicker("Who pays, and for what")
heading("Trust is the line item, not the algorithm")
rows = [
    ("Segment", "Pays for", "Willingness"),
    ("Phage-therapy companies", "Safety screening + auditable reports", "High"),
    ("Biosecurity / gov / public health", "Calibrated, auditable hazard screening", "High · slow"),
    ("Clinical micro / diagnostics", "Validated, reproducible AMR calls + support", "High · regulated"),
    ("Biotech / pharma R&D · CROs", "Hosted API, integration, SLAs, no ops", "Medium"),
    ("Academic labs", "Adoption & citations (not revenue)", "Channel"),
]
table = doc.add_table(rows=len(rows), cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, cell in enumerate(table.rows[0].cells):
    cell.paragraphs[0].runs
    run = cell.paragraphs[0].add_run(rows[0][j])
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for i in range(1, len(rows)):
    for j in range(3):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(rows[i][j])
        run.font.size = Pt(10)
        if j == 0:
            run.bold = True
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 4. Model ─────────────────────────────────────────────────────────────────
kicker("Business model")
heading("Open-core — give away the engine, monetize the layer")
body([("The free tools' own moat is open-source adoption. Match it: open the "
       "engine for credibility and citations, and charge for everything a lab "
       "can't get from a command-line tool on a cluster. GIAE already ships the "
       "account, team, API-key and job-history plumbing this requires.", {})])
bullet([("Free — the engine. ", {"bold": True}), ("Open-source CLI + library. Adoption, trust, citations.", {})])
bullet([("SaaS (usage) — hosted platform. ", {"bold": True}),
        ("API + UI, teams, audit logs, calibrated confidence. Per-genome or "
         "subscription for labs that don't run pipelines.", {})])
bullet([("Enterprise — vertical + on-prem. ", {"bold": True}),
        ("Phage safety product, validation docs, SLAs, on-prem for gov/pharma. "
         "High-ACV contracts.", {})])

# ── 5. Moat ──────────────────────────────────────────────────────────────────
kicker("Why it defends")
heading("The moat compounds; a CLI tool's doesn't")
bullet([("The trust layer. ", {"bold": True}),
        ("Calibration, provenance and abstention are architectural — hard to "
         "retrofit onto a batch CLI. This is the product, and the thing free "
         "tools structurally lack.", {})])
bullet([("Phage-specific depth. ", {"bold": True}),
        ("Dark-gene function, safety knowledge and phage-aware ORF detection "
         "compound into domain expertise competitors would have to rebuild.", {})])
bullet([("A data flywheel. ", {"bold": True}),
        ("Every genome annotated on the platform improves calibration and "
         "dark-gene models. A laptop CLI builds no flywheel; a hosted platform does.", {})])

figure(ROOT / "post_assets" / "recalibration.png",
       "Figure 2 — The trust differentiator, measured. GIAE's raw confidence is "
       "over-confident (grey); a 5-fold cross-validated recalibration (blue) "
       "collapses calibration error from ECE 0.30 to 0.004, turning the score "
       "into a real probability of correctness. This is what a free CLI cannot "
       "offer — and what a regulated buyer will pay for.", width=5.4)

# ── 6. Risks ─────────────────────────────────────────────────────────────────
kicker("The risks — named, not hidden")
heading("What has to be true")
bullet([("The $0 price anchor. ", {"bold": True}),
        ("Free incumbents set the expectation → verticalize on trust and "
         "safety, never compete on “annotation.”", {})])
bullet([("The functional layer is still maturing. ", {"bold": True}),
        ("The calibration story leans on it; near-term paid value is trust + "
         "safety + UX, with product-name annotation improving on the roadmap.", {})])
bullet([("Regulated verticals sell slowly. ", {"bold": True}),
        ("Biosecurity and clinical have long cycles → a self-serve SaaS tier "
         "carries cash flow while enterprise deals bake.", {})])

# ── Closer ───────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(2)
close = doc.add_paragraph()
close.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = close.add_run("Not a faster annotation tool — the trust layer for genome "
                   "interpretation.")
cr.bold = True
cr.font.size = Pt(13)
cr.font.color.rgb = INDIGO
close2 = doc.add_paragraph()
close2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr2 = close2.add_run("Starting with phage therapy, where a wrong call reaches a patient.")
cr2.italic = True
cr2.font.color.rgb = MUTED

foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(18)
fr = foot.add_run("Figures cited (ECE 0.004, Prodigal parity, 35-genome benchmark) "
                  "are engineering results from the GIAE benchmark and calibration "
                  "work, not market projections. No market-size figures are asserted "
                  "without a defensible source.")
fr.italic = True
fr.font.size = Pt(8)
fr.font.color.rgb = MUTED

doc.save(OUT)
print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")
